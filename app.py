import os
from dotenv import load_dotenv
import streamlit as st
import streamlit.components.v1 as components
import google.generativeai as genai
import PyPDF2
import docx
from io import BytesIO
import json
from streamlit_lottie import st_lottie
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.pagesizes import letter
import re
from docx import Document as DocxDocument
from docx.shared import Pt

# Load environment variables
load_dotenv()
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY") 

st.set_page_config(page_title="LinkedIn Job Application Assistant-Career Aligner", layout="wide")

# Particle background effect HTML+CSS+JS
particles_html = """
<style>
  /* Make sure particles canvas is behind content */
  #particles-js {
    position: fixed;
    width: 100%;
    height: 100%;
    z-index: -1;  /* behind everything */
    top: 0;
    left: 0;
    background: linear-gradient(135deg, #FF69B4, #FF1493);
  }
</style>

<div id="particles-js"></div>

<!-- particles.js library -->
<script src="https://cdn.jsdelivr.net/particles.js/2.0.0/particles.min.js"></script>

<script>
particlesJS("particles-js", {
  "particles": {
    "number": {
      "value": 80,
      "density": {
        "enable": true,
        "value_area": 800
      }
    },
    "color": {
      "value": ["#ff69b4", "#ff1493", "#ff85c1"]
    },
    "shape": {
      "type": "circle",
      "stroke": {
        "width": 0
      }
    },
    "opacity": {
      "value": 0.3,
      "random": true,
      "anim": {
        "enable": true,
        "speed": 0.5,
        "opacity_min": 0.1,
        "sync": false
      }
    },
    "size": {
      "value": 3,
      "random": true,
      "anim": {
        "enable": false
      }
    },
    "move": {
      "enable": true,
      "speed": 1,
      "direction": "none",
      "random": true,
      "straight": false,
      "out_mode": "out",
      "bounce": false
    }
  },
  "interactivity": {
    "detect_on": "canvas",
    "events": {
      "onhover": {
        "enable": false
      },
      "onclick": {
        "enable": false
      },
      "resize": true
    }
  },
  "retina_detect": true
});
</script>
"""

components.html(particles_html, height=0, width=0)

# Global CSS with hover effect and responsive footer
st.markdown(
    """
    <style>
  

    /* Hover box styling */
    .hover-box {
        background-color: #1E1E1E;
        padding: 20px;
        border-radius: 10px;
        text-align: center;
        transition: transform 0.3s, background-color 0.3s, box-shadow 0.3s;
        cursor: default;
        color: white;
    }
    .hover-box:hover {
        background-color: #2A2A2A;
        transform: scale(1.05);
        box-shadow: 0 0 12px rgba(255, 105, 180, 0.8);
    }

    /* Button hover effects */
    button[kind="primary"] {
        border: 2px solid #FF69B4 !important;
        transition: all 0.3s ease;
    }
    button[kind="primary"]:hover {
        background-color: #FF69B4 !important;
        color: #1E1E1E !important;
        box-shadow: 0 0 12px rgba(255, 105, 180, 0.8) !important;
        transform: scale(1.05) !important;
    }

    /* Download button hover */
    div.stDownloadButton>button {
        border: 2px solid #FF69B4;
        transition: all 0.3s ease;
    }
    div.stDownloadButton>button:hover {
        background-color: #FF69B4;
        color: #1E1E1E;
        box-shadow: 0 0 12px rgba(255, 105, 180, 0.8);
        transform: scale(1.05);
    }

    /* Footer styles */
    .footer {
        position: fixed;
        left: 0;
        bottom: 0;
        width: 100%;
        background: linear-gradient(90deg, #FF69B4 0%, #FF85C1 50%, #FF69B4 100%);
        color: #1E1E1E;
        text-align: center;
        padding: 12px 0;
        font-size: 15px;
        font-weight: 600;
        font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
        box-shadow: 0 -3px 10px rgba(255, 105, 180, 0.5);
        display: flex;
        justify-content: center;
        align-items: center;
        gap: 20px;
        z-index: 9999;
    }
    .footer strong {
        color: #fff;
        margin-right: 10px;
    }
    .footer-link {
        color: #fff;
        text-decoration: none;
        display: flex;
        align-items: center;
        gap: 6px;
        padding: 6px 10px;
        border-radius: 5px;
        transition: background-color 0.3s, color 0.3s, box-shadow 0.3s;
    }
    .footer-link:hover {
        background-color: #fff;
        color: #FF1493;
        text-decoration: none;
        box-shadow: 0 2px 8px rgba(255, 20, 147, 0.5);
    }
    .footer i {
        font-size: 18px;
    }
    @media (max-width: 600px) {
        .footer {
            flex-direction: column;
            font-size: 14px;
            gap: 12px;
            padding: 15px 10px;
        }
    }
    </style>
    """,
    unsafe_allow_html=True,
)

# Load Lottie animation
@st.cache_data
def load_lottiefile(filepath):
    try:
        with open(filepath, "r") as f:
            return json.load(f)
    except FileNotFoundError:
        return None


lottie = load_lottiefile("animations/robot_working.json")
if lottie:
    st_lottie(lottie, height=200, speed=1)

st.markdown(
    """
    <h1 style="color:#FF69B4; text-align:center;">🤖 LinkedIn Job Application Assistant</h1>
    <h4 style="text-align:center; color:#CCCCCC;">ATS-Optimized Resume Analyzer & Editor</h4>
    <hr style="border:1px solid #FF69B4;">
    """,
    unsafe_allow_html=True,
)

# File readers
def read_pdf(file):
    reader = PyPDF2.PdfReader(file)
    return "\n".join(page.extract_text() for page in reader.pages)


def read_docx(file):
    doc = docx.Document(file)
    return "\n".join([p.text for p in doc.paragraphs])


# Gemini setup
if not GEMINI_API_KEY:
    st.error("❌ Please set your GEMINI_API_KEY in a .env file.")
    st.stop()

try:
    genai.configure(api_key=GEMINI_API_KEY)
    model = genai.GenerativeModel("gemini-1.5-flash")
except Exception as e:
    st.error(f"Gemini Error: {e}")
    st.stop()

# Job description input
st.markdown("<h2 style='color:#00FA9A;'>📋 Step 1: Paste Job Description</h2>", unsafe_allow_html=True)
job_desc = st.text_area(
    "Job Description",
    height=300,
    placeholder="Paste the job description here...",
    help="Paste the complete job description you found on LinkedIn.",
)

# Resume upload
st.markdown("<h2 style='color:#00FA9A;'>📄 Step 2: Upload Your Resume</h2>", unsafe_allow_html=True)
resume_file = st.file_uploader(
    "Upload your resume file (PDF or DOCX)",
    type=["pdf", "docx"],
    help="Upload your current resume. It will not be saved or shared.",
)
resume_text = ""
if resume_file:
    if resume_file.name.endswith(".pdf"):
        resume_text = read_pdf(BytesIO(resume_file.getvalue()))
    elif resume_file.name.endswith(".docx"):
        resume_text = read_docx(BytesIO(resume_file.getvalue()))
    st.success("✅ Resume loaded successfully.")


# Resume analysis
def analyze_resume(job_desc, resume_text):
    prompt = f"""
You are a professional resume analyzer. Respond ONLY with valid JSON.
Structure:
{{
  "overall_match_percentage": "85%",
  "matching_skills": [{{"skill_name": "Python", "is_match": true}}],
  "missing_skills": [{{"skill_name": "Docker", "is_match": false, "suggestion": "Get Docker certified"}}],
  "experience_match_analysis": "...",
  "education_match_analysis": "...",
  "recommendations_for_improvement": [{{"recommendation": "Add metrics", "section": "Experience", "guidance": "Use quantifiable data"}}],
  "ats_optimization_suggestions": [{{"section": "Skills", "suggested_change": "Add ReactJS", "keywords_to_add": ["ReactJS"], "formatting_suggestion": "Use bullet points", "reason": "Increases ATS visibility"}}],
  "key_strengths": "...",
  "areas_of_improvement": "..."
}}

Job Description:
{job_desc}

Resume:
{resume_text}
"""
    try:
        res = model.generate_content(prompt)
        response_text = res.text.strip().strip("` ")
        try:
            return json.loads(response_text)
        except json.JSONDecodeError:
            import re

            match = re.search(r"\{.*\}", response_text, re.DOTALL)
            if match:
                try:
                    return json.loads(match.group(0))
                except Exception:
                    pass
            st.error("❌ Gemini response is not valid JSON. Full response shown below:")
            st.code(response_text)
            return {}
    except Exception as e:
        st.error(f"❌ Error analyzing resume: {e}")
        return {}


if resume_file and job_desc:
    with st.spinner("Analyzing resume with Gemini..."):
        result = analyze_resume(job_desc, resume_text)

    if result:
        st.subheader("📊 Resume Analysis Dashboard")

        col1, col2, col3 = st.columns(3)
        with col1:
            st.markdown(
                f"""
                <div class="hover-box">
                    <h3 style="color:#FF69B4;">🎯 Match %</h3>
                    <h1 style="color:#FFFFFF;">{result.get("overall_match_percentage", "0%")}</h1>
                </div>
                """,
                unsafe_allow_html=True,
            )
        with col2:
            st.markdown(
                f"""
                <div class="hover-box">
                    <h3 style="color:#00FA9A;">✅ Matching Skills</h3>
                    <h1 style="color:#FFFFFF;">{len(result.get("matching_skills", []))}</h1>
                </div>
                """,
                unsafe_allow_html=True,
            )
        with col3:
            st.markdown(
                f"""
                <div class="hover-box">
                    <h3 style="color:#FFD700;">⚠️ Missing Skills</h3>
                    <h1 style="color:#FFFFFF;">{len(result.get("missing_skills", []))}</h1>
                </div>
                """,
                unsafe_allow_html=True,
            )

        tabs = st.tabs(
            [
                "🟢 Skills Analysis",
                "📁 Experience & Education",
                "💡 Recommendations",
                "💌 Cover Letter",
                "📄 Updated Resume",
            ]
        )

        with tabs[0]:
            st.markdown("<h3 style='color:#FF69B4;'>✅ Matching Skills</h3>", unsafe_allow_html=True)
            for skill in result.get("matching_skills", []):
                st.success(f"✔ {skill.get('skill_name', '')}")
            st.markdown("<h3 style='color:#FFD700;'>⚠️ Missing Skills</h3>", unsafe_allow_html=True)
            for skill in result.get("missing_skills", []):
                st.warning(f"✖ {skill.get('skill_name', '')}")
                st.info(f"Suggestion: {skill.get('suggestion', '')}")

        with tabs[1]:
            st.markdown("<h3 style='color:#FF69B4;'>📁 Experience Match</h3>", unsafe_allow_html=True)
            st.write(result.get("experience_match_analysis", ""))
            st.markdown("<h3 style='color:#FF69B4;'>🎓 Education Match</h3>", unsafe_allow_html=True)
            st.write(result.get("education_match_analysis", ""))

        with tabs[2]:
            st.markdown("<h3 style='color:#FF69B4;'>💡 Recommendations</h3>", unsafe_allow_html=True)
            for rec in result.get("recommendations_for_improvement", []):
                st.info(f"{rec.get('recommendation', '')} — {rec.get('guidance', '')} in {rec.get('section', '')}")
            st.markdown("<h3 style='color:#FF69B4;'>🤖 ATS Suggestions</h3>", unsafe_allow_html=True)
            for rec in result.get("ats_optimization_suggestions", []):
                st.warning(f"Section: {rec.get('section', '')}")
                st.write(f"Suggested: {rec.get('suggested_change', '')}")
                if rec.get("keywords_to_add"):
                    st.write(f"Keywords: {', '.join(rec.get('keywords_to_add'))}")
                if rec.get("reason"):
                    st.write(f"Reason: {rec.get('reason')}")

        with tabs[3]:
            st.markdown("<h3 style='color:#FF69B4;'>💌 Cover Letter</h3>", unsafe_allow_html=True)
            if st.button("🚀 Generate Cover Letter"):
                prompt = f"""
Write a professional, ATS-friendly cover letter using:
Resume: {resume_text}
Job: {job_desc}
Match Data: {json.dumps(result, indent=2)}
"""
                try:
                    cover = model.generate_content(prompt).text
                    st.text_area("📄 Generated Cover Letter", cover, height=400)
                    st.download_button("📥 Download Cover Letter", cover, "cover_letter.txt")
                except Exception as e:
                    st.error(f"Error: {e}")

        with tabs[4]:
            st.markdown("<h3 style='color:#FF69B4;'>📄 Updated ATS-Friendly Resume</h3>", unsafe_allow_html=True)

            def generate_updated_resume(text, suggestions):
                buffer = BytesIO()
                doc = SimpleDocTemplate(
                    buffer,
                    pagesize=letter,
                    rightMargin=40,
                    leftMargin=40,
                    topMargin=40,
                    bottomMargin=30,
                )
                styles = getSampleStyleSheet()
                heading_style = ParagraphStyle(
                    name="Heading",
                    parent=styles["Heading2"],
                    fontName="Helvetica-Bold",
                    fontSize=14,
                    spaceAfter=10,
                )
                normal_style = ParagraphStyle(name="Normal", fontName="Helvetica", fontSize=10, leading=12)
                bullet_style = ParagraphStyle(parent=normal_style, name="Bullet", bulletIndent=0, leftIndent=10)
                content = []
                content.append(Paragraph("Updated ATS-Friendly Resume", heading_style))
                content.append(Spacer(1, 12))
                sections = re.split(r"\n\s*(?=[A-Z][a-z]+:)", text)
                for sec in sections:
                    lines = sec.strip().split("\n")
                    for line in lines:
                        content.append(
                            Paragraph(line.strip(), bullet_style if line.strip().startswith(("-", "•")) else normal_style)
                        )
                    content.append(Spacer(1, 8))
                if suggestions:
                    content.append(Spacer(1, 16))
                    content.append(Paragraph("ATS Optimization Suggestions", heading_style))
                    for s in suggestions:
                        content.append(Paragraph(f"<b>Section:</b> {s.get('section', '')}", normal_style))
                        content.append(Paragraph(f"<b>Suggested Change:</b> {s.get('suggested_change', '')}", normal_style))
                        keywords = ", ".join(s.get("keywords_to_add", []))
                        if keywords:
                            content.append(Paragraph(f"<b>Keywords to Add:</b> {keywords}", normal_style))
                        if s.get("reason"):
                            content.append(Paragraph(f"<b>Reason:</b> {s.get('reason')}", normal_style))
                        content.append(Spacer(1, 8))
                doc.build(content)
                buffer.seek(0)
                return buffer

            def generate_docx_resume(text, suggestions):
                doc = DocxDocument()
                doc.add_heading("Updated ATS-Friendly Resume", level=1)
                sections = re.split(r"\n\s*(?=[A-Z][a-z]+:)", text)
                for sec in sections:
                    lines = sec.strip().split("\n")
                    for line in lines:
                        clean_line = re.sub(r"[\x00-\x1F\x7F]", "", line.strip())
                        if clean_line:
                            para = doc.add_paragraph()
                            run = para.add_run(clean_line)
                            run.font.size = Pt(10)
                if suggestions:
                    doc.add_page_break()
                    doc.add_heading("ATS Optimization Suggestions", level=2)
                    for s in suggestions:
                        doc.add_paragraph(f"Section: {s.get('section', '')}", style="List Bullet")
                        doc.add_paragraph(f"Suggested Change: {s.get('suggested_change', '')}")
                        keywords = ", ".join(s.get("keywords_to_add", []))
                        if keywords:
                            doc.add_paragraph(f"Keywords to Add: {keywords}")
                        if s.get("reason"):
                            doc.add_paragraph(f"Reason: {s.get('reason')}")
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                return buffer

            updated_pdf = generate_updated_resume(resume_text, result.get("ats_optimization_suggestions", []))
            updated_docx = generate_docx_resume(resume_text, result.get("ats_optimization_suggestions", []))

            st.download_button(
                "📥 Download Updated Resume (PDF)",
                updated_pdf,
                "updated_resume.pdf",
                mime="application/pdf",
            )
            st.download_button(
                "📥 Download Updated Resume (DOCX)",
                updated_docx,
                "updated_resume.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

# Footer HTML with icons and hover effect
st.markdown(
    """
    <!-- Font Awesome CDN -->
    <link
      href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.4.0/css/all.min.css"
      rel="stylesheet"
    />
    <div class="footer">
      <strong>Made with ❤️ by Shrikrishnasutar.dev</strong>
      <a href="https://github.com/Skrishna0703" target="_blank" class="footer-link">
        <i class="fab fa-github"></i> GitHub
      </a>
      <a href="https://www.linkedin.com/in/shrikrishna-sutar-3b601524b" target="_blank" class="footer-link">
        <i class="fab fa-linkedin"></i> LinkedIn
      </a>
      <a href="mailto:shrikrishna0703@gmail.com" class="footer-link">
        <i class="fas fa-envelope"></i> Contact
      </a>
    </div>
    """,
    unsafe_allow_html=True,
)
